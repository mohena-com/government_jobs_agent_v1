from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def link(p, text, url):
    if not url: return
    rid=p.part.relate_to(url,"http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",is_external=True)
    h=OxmlElement("w:hyperlink"); h.set(qn("r:id"),rid)
    r=OxmlElement("w:r"); t=OxmlElement("w:t"); t.text=text
    r.append(t); h.append(r); p._p.append(h)

def build(rows,out,today):
    doc=Document(); doc.styles["Normal"].font.name="Aptos"; doc.styles["Normal"].font.size=Pt(9)
    doc.add_heading(f"SarkariResult Future Jobs — {today:%d %B %Y}",0)
    doc.add_paragraph("Discovery source: SarkariResult Latest Jobs. Only listings with a parsed Last Date strictly later than the report date are included. SarkariResult is not treated as the final authority.")
    table=doc.add_table(rows=1,cols=5)
    for c,t in zip(table.rows[0].cells,["#","Job","Last Date","Extended","Detail"]): c.text=t
    for i,r in enumerate(rows,1):
        c=table.add_row().cells
        c[0].text=str(i); c[1].text=r["title"]; c[2].text=r["last_date"].strftime("%d/%m/%Y"); c[3].text="Yes" if r.get("date_extended") else "No"
        link(c[4].paragraphs[0],"Open",r["url"])
    Path(out).parent.mkdir(parents=True,exist_ok=True); doc.save(out)
