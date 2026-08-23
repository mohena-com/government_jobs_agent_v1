
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


SOCIAL_GARBAGE = {
    "telegram", "join us", "whatsapp", "instagram",
    "follow", "x", "image",
}


def clean_report_value(value):
    if value is None:
        return ""

    cleaned_lines = []

    for raw_line in str(value).splitlines():
        line = raw_line.strip()
        if not line:
            continue

        line = line.lstrip("·•*- ").strip()
        if not line:
            continue

        low = line.lower()

        if low in SOCIAL_GARBAGE:
            continue

        if (
            ("telegram" in low or "whatsapp" in low or "instagram" in low)
            and len(line) < 100
        ):
            continue

        if (
            (low.startswith("join us") or low.startswith("follow"))
            and len(line) < 100
        ):
            continue

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines).strip()


def add_hyperlink(paragraph, text, url):
    if not url:
        return

    rid = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    link = OxmlElement("w:hyperlink")
    link.set(qn("r:id"), rid)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    rPr.append(underline)

    run.append(rPr)

    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)

    link.append(run)
    paragraph._p.append(link)


def set_cell_text(cell, text, bold=False):
    cell.text = ""

    cleaned = clean_report_value(text) or "Not found"

    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)

    r = p.add_run(cleaned)
    r.bold = bold
    r.font.size = Pt(8.5)

    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tcPr.append(shd)


def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")

    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))

        if element is None:
            element = OxmlElement(tag)
            borders.append(element)

        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "B7B7B7")


def add_section_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    r = p.add_run(text.upper())
    r.bold = True
    r.font.size = Pt(10)

    return p


def add_key_value_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    for label, value in rows:
        cells = table.add_row().cells

        set_cell_text(cells[0], label, True)
        set_cell_text(cells[1], value or "Not found")
        set_cell_shading(cells[0], "EDEDED")

    set_table_borders(table)
    return table


def _clean_section_lines(text):
    if not text:
        return []

    lines = []

    for raw in str(text).splitlines():
        line = raw.strip()
        if not line:
            continue

        line = line.lstrip("·•*- ").strip()
        line = clean_report_value(line)

        if line:
            lines.append(line)

    return lines


def _compact_label_value_lines(lines):
    result = []
    i = 0

    while i < len(lines):
        current = lines[i].strip()

        if current.endswith(":") and i + 1 < len(lines):
            nxt = lines[i + 1].strip()

            if nxt and not nxt.endswith(":"):
                result.append(
                    current.rstrip(":").strip() + ": " + nxt
                )
                i += 2
                continue

        result.append(current)
        i += 1

    return result


def add_compact_section(doc, text, mode="lines"):
    lines = _clean_section_lines(text)

    if not lines:
        doc.add_paragraph("Not found")
        return

    for line in _compact_label_value_lines(lines):
        line = (
            line.replace("**", "")
                .replace("__", "")
                .strip()
        )

        line = clean_report_value(line)
        if not line:
            continue

        if mode == "eligibility":
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.keep_together = True

            if re.match(r"^Management Trainee\s*\(", line, re.I):
                r = p.add_run(line)
                r.bold = True
                r.font.size = Pt(9)
            else:
                p.paragraph_format.left_indent = Inches(0.18)
                p.add_run(line)

        elif mode == "bullets":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.left_indent = Inches(0.18)
            p.add_run(line)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.keep_together = True
            p.add_run(line)


def add_vacancy_table(doc, rows):
    if not rows:
        doc.add_paragraph("No structured vacancy table detected.")
        return

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header = table.rows[0]
    set_repeat_table_header(header)

    set_cell_text(header.cells[0], "Post", True)
    set_cell_text(header.cells[1], "Vacancies", True)

    set_cell_shading(header.cells[0], "D9EAF7")
    set_cell_shading(header.cells[1], "D9EAF7")

    for row in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], row.get("post_name", ""))
        set_cell_text(cells[1], row.get("vacancies", ""))

    set_table_borders(table)


def add_links(doc, job):
    links = []

    for link in job.get("application_links", []):
        url = link.get("url")
        if url:
            links.append(("Apply Online", url))

    for link in job.get("notification_links", []):
        url = link.get("url")
        if url:
            links.append(("Official Notification", url))

    unique = []
    seen = set()

    for label, url in links:
        key = (label, url)
        if key not in seen:
            seen.add(key)
            unique.append((label, url))

    existing_urls = {url for _, url in unique}

    for link in job.get("official_candidates", []):
        url = link.get("url")
        if url and url not in existing_urls:
            unique.append(
                (
                    link.get("text") or "Official Website",
                    url,
                )
            )
            existing_urls.add(url)

    if not unique:
        doc.add_paragraph(
            "No external application/notification link detected."
        )
        return

    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    header = table.rows[0]
    set_repeat_table_header(header)

    set_cell_text(header.cells[0], "Type", True)
    set_cell_text(header.cells[1], "URL", True)

    set_cell_shading(header.cells[0], "D9EAF7")
    set_cell_shading(header.cells[1], "D9EAF7")

    for label, url in unique:
        cells = table.add_row().cells

        set_cell_text(cells[0], label, True)

        cells[1].text = ""
        p = cells[1].paragraphs[0]
        p.paragraph_format.space_after = Pt(0)

        # Display the actual URL, not "Open link".
        add_hyperlink(p, url, url)

    set_table_borders(table)


def _configure_document(doc):
    section = doc.sections[0]
    section.top_margin = Inches(0.45)
    section.bottom_margin = Inches(0.45)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    normal = doc.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(8.5)


def add_job(doc, job, index=None):
    listing = job.get("listing", {})

    title = (
        job.get("post_title")
        or listing.get("title")
        or "Recruitment Notice"
    )

    if index is not None and index > 1:
        doc.add_page_break()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(15)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    organisation = clean_report_value(
        job.get("organisation")
    ) or "Organisation not identified"

    r = p.add_run(organisation)
    r.bold = True
    r.font.size = Pt(10)

    deadline = (
        job.get("application_end")
        or listing.get("last_date")
        or "Not found"
    )

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run(
        "APPLICATION DEADLINE: "
        + clean_report_value(deadline)
    )
    r.bold = True
    r.font.size = Pt(10)

    add_section_heading(doc, "Key Information")

    add_key_value_table(
        doc,
        [
            (
                "Advertisement / Reference No.",
                job.get("advertisement_number"),
            ),
            (
                "Published / Updated",
                job.get("post_update"),
            ),
            (
                "Total Vacancies",
                job.get("total_vacancies"),
            ),
            (
                "Application Start",
                job.get("application_start"),
            ),
            (
                "Application End",
                job.get("application_end"),
            ),
            (
                "Age Limit",
                job.get("age_limit"),
            ),
        ],
    )

    add_section_heading(doc, "Important Dates")
    add_compact_section(
        doc,
        job.get("important_dates_raw"),
        mode="compact",
    )

    add_section_heading(doc, "Application Fee")
    add_compact_section(
        doc,
        job.get("application_fee"),
        mode="compact",
    )

    add_section_heading(doc, "Vacancy Details")
    add_vacancy_table(
        doc,
        job.get("vacancy_rows", []),
    )

    add_section_heading(doc, "Eligibility")
    add_compact_section(
        doc,
        job.get("eligibility"),
        mode="eligibility",
    )

    add_section_heading(doc, "Selection Process")
    add_compact_section(
        doc,
        job.get("selection_process"),
        mode="bullets",
    )

    add_section_heading(doc, "Pay / Salary")
    add_compact_section(
        doc,
        job.get("pay_scale"),
        mode="bullets",
    )

    add_section_heading(doc, "How to Apply")
    add_compact_section(
        doc,
        job.get("how_to_apply"),
        mode="bullets",
    )

    add_section_heading(doc, "Official Links")
    add_links(doc, job)

    add_section_heading(doc, "Source")

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    p.add_run("SarkariResult detail page: ").bold = True

    add_hyperlink(
        p,
        job.get("detail_url", ""),
        job.get("detail_url", ""),
    )


def _safe_filename(text, max_len=90):
    text = clean_report_value(text) or "Recruitment"
    text = re.sub(r"[^\w\s.-]", "", text)
    text = re.sub(r"\s+", "_", text).strip("_")
    return text[:max_len] or "Recruitment"


def make_summary_report(today, results, out):
    """
    Create ONLY the summary/index DOCX.

    The summary contains the report title and one index table.
    Detailed recruitment information is written to separate DOCX files.
    """
    doc = Document()
    _configure_document(doc)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run(
        f"Government Jobs Report — {today:%d %B %Y}"
    )
    r.bold = True
    r.font.size = Pt(18)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r = p.add_run(
        "Source: SarkariResult Latest Jobs | "
        "Future last-date listings only"
    )
    r.font.size = Pt(9)

    add_section_heading(doc, "Index")

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    header = table.rows[0]
    set_repeat_table_header(header)

    for cell, label in zip(
        header.cells,
        ["#", "Organisation", "Post", "Last Date"],
    ):
        set_cell_text(cell, label, True)
        set_cell_shading(cell, "D9EAF7")

    for i, job in enumerate(results, 1):
        listing = job.get("listing", {})

        cells = table.add_row().cells

        set_cell_text(cells[0], i)
        set_cell_text(
            cells[1],
            job.get("organisation") or "Not identified",
        )
        set_cell_text(
            cells[2],
            job.get("post_title")
            or listing.get("title", ""),
        )
        set_cell_text(
            cells[3],
            job.get("application_end")
            or listing.get("last_date", ""),
        )

    set_table_borders(table)

    Path(out).parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    doc.save(out)

    return out


def make_job_reports(today, results, output_dir):
    """
    Create one DOCX per recruitment.
    """
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    files = []

    for i, job in enumerate(results, 1):
        listing = job.get("listing", {})

        title = (
            job.get("post_title")
            or listing.get("title")
            or f"Recruitment_{i}"
        )

        filename = (
            f"{i:02d}_"
            f"{_safe_filename(title)}_"
            f"{today.isoformat()}.docx"
        )

        path = output_dir / filename

        doc = Document()
        _configure_document(doc)
        add_job(doc, job)

        doc.save(path)
        files.append(path)

    return files


def make_report(today, results, out):
    """
    V1.4 output:

        reports/
        ├── SarkariResult_LatestJobs_YYYY-MM-DD_Summary.docx
        └── jobs/
            ├── 01_....docx
            ├── 02_....docx
            └── ...

    Returns:
        (summary_path, [job_paths...])
    """
    out = Path(out)

    # Keep the user-facing summary filename predictable.
    summary_out = out.with_name(
        out.stem + "_Summary" + out.suffix
    )

    jobs_dir = out.parent / "jobs"

    summary_path = make_summary_report(
        today,
        results,
        summary_out,
    )

    job_files = make_job_reports(
        today,
        results,
        jobs_dir,
    )

    return summary_path, job_files
