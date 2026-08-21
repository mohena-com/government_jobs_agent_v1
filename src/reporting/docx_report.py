from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import date
def link(p,text,url):
    if not url:return
    rid=p.part.relate_to(url,'http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink',is_external=True); h=OxmlElement('w:hyperlink'); h.set(qn('r:id'),rid); rr=OxmlElement('w:r'); rp=OxmlElement('w:rPr'); u=OxmlElement('w:u'); u.set(qn('w:val'),'single'); rp.append(u); rr.append(rp); t=OxmlElement('w:t'); t.text=text; rr.append(t); h.append(rr); p._p.append(h)
def field(d,label,val): p=d.add_paragraph(); p.add_run(label+': ').bold=True; p.add_run(str(val or 'Not stated'))
def build_report(records,path):
    d=Document(); d.styles['Normal'].font.name='Aptos'; d.styles['Normal'].font.size=Pt(9); d.add_heading('Government Jobs Report — '+date.today().strftime('%d %B %Y'),0); d.add_paragraph('New and updated recruitment notifications identified by Government Jobs Agent V1.')
    t=d.add_table(rows=1,cols=5)
    for i,x in enumerate(['No.','Organisation','Post','Last date','Status']):t.rows[0].cells[i].text=x
    for i,r in enumerate(records,1):
        c=t.add_row().cells;c[0].text=str(i);c[1].text=r.get('organisation') or '—';c[2].text=r.get('post_title') or '—';c[3].text=str(r.get('application_end_date') or '—');c[4].text='VERIFIED' if r.get('source_verified') else 'UNVERIFIED'
    for r in records:
        d.add_page_break(); d.add_heading(r.get('post_title') or 'Recruitment Notice',1)
        for label,key in [('Organisation','organisation'),('Ministry / Department','ministry_department'),('Post','post_title'),('Vacancies','vacancies_total'),('Advertisement number','advertisement_number'),('Notification number','notification_number'),('Published / updated','publication_date'),('Application period','application_start_date')]:field(d,label,r.get(key))
        field(d,'Application closing date',r.get('application_end_date')); field(d,'Eligibility',r.get('qualification')); field(d,'Experience',r.get('experience')); field(d,'Age limit',r.get('age_limit')); field(d,'Age relaxation',r.get('age_relaxation')); field(d,'Pay',r.get('pay_level') or r.get('pay_scale') or r.get('salary')); field(d,'Application fee',r.get('application_fee')); field(d,'Selection process',r.get('selection_process')); field(d,'Job location',r.get('job_location'))
        p=d.add_paragraph();p.add_run('Application website: ').bold=True;link(p,'Apply Online',r.get('application_url'))
        p=d.add_paragraph();p.add_run('Official notification: ').bold=True;link(p,'View Notification',r.get('notification_url'))
        field(d,'Source verification','Official government domain verified.' if r.get('source_verified') else 'Could not verify official government domain.');field(d,'Extraction confidence',f"{float(r.get('extraction_confidence') or 0):.0%}")
    Path(path).parent.mkdir(parents=True,exist_ok=True);d.save(path);return path
