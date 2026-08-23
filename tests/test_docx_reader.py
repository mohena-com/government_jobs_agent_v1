from docx import Document
from src.docx.reader import read_docx, to_locked_facts


def test_reader_extracts_job(tmp_path):
    p = tmp_path / "job.docx"
    doc = Document()
    doc.add_paragraph("RCFL Management Trainee Recruitment 2026 Apply Online for 94 Post")
    t = doc.add_table(rows=3, cols=2)
    t.cell(0, 0).text = "Name Of Post :"
    t.cell(0, 1).text = "RCFL Management Trainee Recruitment 2026 Apply Online for 94 Post"
    t.cell(1, 0).text = "Total Vacancies"
    t.cell(1, 1).text = "94"
    t.cell(2, 0).text = "Application End"
    t.cell(2, 1).text = "24/08/2026"
    v = doc.add_table(rows=3, cols=2)
    v.cell(0, 0).text = "Post Name"
    v.cell(0, 1).text = "Total Post"
    v.cell(1, 0).text = "Management Trainee (Chemical)"
    v.cell(1, 1).text = "32"
    v.cell(2, 0).text = "Management Trainee (IT)"
    v.cell(2, 1).text = "05"
    doc.save(p)

    parsed = read_docx(p)
    assert parsed["job_count"] == 1
    job = parsed["jobs"][0]
    assert job["title"].startswith("RCFL Management Trainee")
    assert job["fields"]["total_vacancies"] == "94"
    assert len(job["vacancy_rows"]) == 2

    facts = to_locked_facts(job)
    assert facts["total_vacancies"] == "94"
