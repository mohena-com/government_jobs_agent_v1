from docx import Document
from src.docx.reader import read_docx, to_locked_facts
from src.docx.quality_gate import quality_gate


def test_v16_rvunl_layout_extracts_key_information(tmp_path):
    p = tmp_path / "rvunl.docx"
    doc = Document()
    doc.add_paragraph("Rajasthan RVUNL for JE, Junior Accountant & Junior Assistant/ Commercial Assistant-II Common Recruitment 2026 Apply Online for 2005 Post")
    doc.add_paragraph("Rajasthan Rajya Vidyut Utpadan Nigam Ltd. (RVUNL)")
    doc.add_paragraph("APPLICATION DEADLINE: 25 August 2026")
    doc.add_paragraph("Key Information")
    t = doc.add_table(rows=7, cols=2)
    rows = [
        ("Advertisement / Reference No.", "BULUN/Rectt 3036-27/01"),
        ("Published / Updated", "05 August 2026 | 11:13 PM"),
        ("Total Vacancies", "2005"),
        ("Application Start", "10 August 2026"),
        ("Application End", "25 August 2026"),
        ("Age Limit", "18-40 years"),
        ("", ""),
    ]
    for i, (a, b) in enumerate(rows):
        t.cell(i, 0).text = a
        t.cell(i, 1).text = b
    doc.add_paragraph("Vacancy Details")
    v = doc.add_table(rows=3, cols=2)
    v.cell(0,0).text="Post"
    v.cell(0,1).text="Vacancies"
    v.cell(1,0).text="Junior Engineer"
    v.cell(1,1).text="1500"
    v.cell(2,0).text="Junior Accountant"
    v.cell(2,1).text="505"
    doc.add_paragraph("Eligibility")
    doc.add_paragraph("Post-wise qualification as stated in the official notification.")
    doc.add_paragraph("Official Links")
    l = doc.add_table(rows=2, cols=2)
    l.cell(0,0).text="Type"; l.cell(0,1).text="URL"
    l.cell(1,0).text="Official Notification"; l.cell(1,1).text="https://jankalyanfile.rajasthan.gov.in/test.pdf"
    doc.save(p)

    parsed = read_docx(p)
    assert parsed["job_count"] == 1
    job = parsed["jobs"][0]
    facts = to_locked_facts(job)
    assert facts["organisation"] == "Rajasthan Rajya Vidyut Utpadan Nigam Ltd. (RVUNL)"
    assert facts["advertisement_number"] == "BULUN/Rectt 3036-27/01"
    assert facts["total_vacancies"] == "2005"
    assert facts["application_start"] == "10 August 2026"
    assert facts["application_end"] == "25 August 2026"
    assert facts["age_limit"] == "18-40 years"
    assert facts["derived_vacancy_sum"] == 2005
    # V1.9.8 hard gate requires verified post-specific eligibility. The DOCX
    # fixture itself does not contain it, so provide the verified canonical
    # post facts as the official-verification layer would.
    facts["post_facts"] = [
        {"post": "Junior Engineer-I (Electrical)", "qualification": "Verified"},
        {"post": "Junior Engineer-I (Mechanical)", "qualification": "Verified"},
        {"post": "Junior Engineer-I (Civil)", "qualification": "Verified"},
        {"post": "Junior Accountant", "qualification": "Verified"},
        {"post": "Junior Assistant/ Commercial Assistant-II", "qualification": "Verified"},
    ]
    assert quality_gate(job, facts)["status"] == "PASS"
